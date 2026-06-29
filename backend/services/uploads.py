from io import BytesIO
from pathlib import Path

from config import get_settings
from schemas import DocumentCreate


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


class DocumentUploadError(ValueError):
    pass


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentUploadError("Could not decode text file")


def _extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise DocumentUploadError("PDF support is not installed") from exc

    try:
        reader = PdfReader(BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(page.strip() for page in pages if page.strip())
    except Exception as exc:
        raise DocumentUploadError("Could not extract text from PDF") from exc


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document as DocxDocument
    except Exception as exc:
        raise DocumentUploadError("Word (.docx) support is not installed") from exc

    try:
        document = DocxDocument(BytesIO(data))
        blocks: list[str] = [para.text for para in document.paragraphs if para.text.strip()]
        # Pull text out of tables too — Word documents often carry data in tables.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        return "\n\n".join(blocks)
    except DocumentUploadError:
        raise
    except Exception as exc:
        raise DocumentUploadError("Could not extract text from Word document") from exc


def document_from_upload(filename: str, data: bytes, title: str | None = None) -> DocumentCreate:
    settings = get_settings()
    if not filename:
        raise DocumentUploadError("Uploaded file must include a filename")
    if not data:
        raise DocumentUploadError("Uploaded file is empty")
    if len(data) > settings.max_upload_bytes:
        raise DocumentUploadError("Uploaded file is too large")

    suffix = Path(filename).suffix.lower()
    if suffix == ".doc":
        raise DocumentUploadError("Legacy .doc files are not supported — save as .docx, .pdf, or .txt")
    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentUploadError("Unsupported file type. Use .txt, .md, .pdf, or .docx")

    if suffix == ".pdf":
        content = _extract_pdf_text(data)
        source_type = "pdf"
    elif suffix == ".docx":
        content = _extract_docx_text(data)
        source_type = "docx"
    else:
        content = _decode_text(data)
        source_type = "document"

    content = content.strip()
    if len(content) > settings.max_document_chars:
        content = content[: settings.max_document_chars].rstrip()
    if len(content) < 20:
        raise DocumentUploadError("Uploaded file does not contain enough readable text")

    document_title = (title or Path(filename).stem).strip()
    return DocumentCreate(title=document_title, content=content, source_type=source_type)
